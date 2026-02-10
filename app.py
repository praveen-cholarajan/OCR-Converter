import os
import sys
import time
import subprocess
import threading
import multiprocessing
import logging
from flask import Flask, request, render_template, send_file, flash, redirect, url_for, session, jsonify
from werkzeug.utils import secure_filename
import pytesseract
from docx import Document
from PIL import Image
import uuid
from threading import Thread
from concurrent.futures import ProcessPoolExecutor, as_completed
import shutil
import tempfile
import re
import hashlib
from datetime import datetime, timedelta
from typing import Optional, Tuple, Dict, Any
from pathlib import Path

# Tesseract executable path (can be overridden with TESSERACT_CMD env var)
TESSERACT_CMD = os.environ.get('TESSERACT_CMD', r"C:\Program Files\Tesseract-OCR\tesseract.exe")
try:
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD
except Exception:
    pass

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("app.log"),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)
try:
    logger.info(f"Tesseract command set to: {TESSERACT_CMD}")
except Exception:
    pass

app = Flask(__name__)

# Configuration
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'pdf'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 64 * 1024 * 1024  # 64 MB limit
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(hours=1)  # Session timeout

# Use a strong secret key from environment or generate one
app.secret_key = os.environ.get('SECRET_KEY', hashlib.sha256(os.urandom(32)).hexdigest())

# Docker-specific configuration
DOCKER_ENV = os.environ.get('DOCKER_ENV', 'false').lower() == 'true'

# Ensure the upload folder exists
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

# Add storage for background tasks
TASK_STATUS = {}  # Store task status updates
TASK_RESULTS = {}  # Store task results

# Setup periodic cleanup
CLEANUP_INTERVAL = 3600  # 1 hour in seconds
TASK_TIMEOUT = 3600  # 1 hour in seconds
LAST_CLEANUP_TIME = time.time()

def allowed_file(filename: Optional[str]) -> bool:
    """Check if a file extension is allowed."""
    if not filename:
        return False
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def secure_clean_filename(filename: str) -> str:
    """Secure and sanitize the filename."""
    filename = secure_filename(filename)
    filename = re.sub(r'[^\w\s.-]', '', filename)
    filename = filename.replace(' ', '_')
    return filename

def cleanup_old_files() -> None:
    """Remove old files from the uploads directory and expired tasks."""
    global LAST_CLEANUP_TIME
    current_time = time.time()
    if current_time - LAST_CLEANUP_TIME < CLEANUP_INTERVAL:
        return
    LAST_CLEANUP_TIME = current_time
    logger.info("Running periodic cleanup")
    try:
        upload_path = Path(UPLOAD_FOLDER)
        for file_path in upload_path.iterdir():
            if file_path.is_file() and current_time - file_path.stat().st_mtime > 86400:
                try:
                    file_path.unlink()
                    logger.info(f"Deleted old file: {file_path}")
                except Exception as e:
                    logger.error(f"Error deleting old file {file_path}: {e}")
    except Exception as e:
        logger.error(f"Error during file cleanup: {e}")
    expired_tasks = []
    for task_id, status in TASK_STATUS.items():
        if status.get("timestamp", 0) + TASK_TIMEOUT < current_time:
            expired_tasks.append(task_id)
    for task_id in expired_tasks:
        TASK_STATUS.pop(task_id, None)
        TASK_RESULTS.pop(task_id, None)
        logger.info(f"Removed expired task: {task_id}")

_dependency_check_cache: Dict[str, Tuple[float, Tuple[bool, str]]] = {}

def check_dependencies() -> Tuple[bool, str]:
    """Check if all required dependencies are installed. Caches result for 60s."""
    cache_key = 'all'
    now = time.time()
    if cache_key in _dependency_check_cache:
        ts, result = _dependency_check_cache[cache_key]
        if now - ts < 60:
            return result
    try:
        if DOCKER_ENV:
            result = (True, "Running in Docker, dependencies assumed to be installed")
            _dependency_check_cache[cache_key] = (now, result)
            return result
        if sys.platform == 'darwin':
            try:
                from pdf2image import convert_from_path
                convert_from_path.get_page_count('test.pdf')
            except Exception as e:
                if "Unable to get page count. Is poppler installed and in PATH?" in str(e):
                    result = (False, "Poppler is not installed or not in PATH. Install it with 'brew install poppler'")
                    _dependency_check_cache[cache_key] = (now, result)
                    return result
        try:
            # Try the configured TESSERACT_CMD first, then fallback to system 'tesseract'
            try:
                subprocess.check_output([TESSERACT_CMD, '--version'], stderr=subprocess.STDOUT)
            except (subprocess.CalledProcessError, FileNotFoundError, OSError):
                subprocess.check_output(['tesseract', '--version'], stderr=subprocess.STDOUT)
        except (subprocess.CalledProcessError, FileNotFoundError, OSError):
            result = (False, "Tesseract OCR is not installed or not in PATH. On macOS, install it with 'brew install tesseract'")
            _dependency_check_cache[cache_key] = (now, result)
            return result
        result = (True, "All dependencies are properly installed")
        _dependency_check_cache[cache_key] = (now, result)
        return result
    except Exception as e:
        result = (False, f"Error checking dependencies: {str(e)}")
        _dependency_check_cache[cache_key] = (now, result)
        return result

def check_dependency(name: str) -> Tuple[bool, Dict[str, Any]]:
    """Check a specific dependency and return detailed information"""
    if DOCKER_ENV:
        return True, {"installed": True, "version": "Docker Environment", "message": "Running in Docker container"}
    
    try:
        if name.lower() == 'poppler':
            try:
                output = subprocess.check_output(['pdftoppm', '-v'], stderr=subprocess.STDOUT, text=True)
                version = output.strip() if output else "Unknown version"
                return True, {"installed": True, "version": version, "message": "Poppler is installed"}
            except (subprocess.CalledProcessError, FileNotFoundError):
                return False, {"installed": False, "message": "Poppler is not installed or not in PATH"}
        
        elif name.lower() == 'tesseract':
            try:
                # Prefer configured TESSERACT_CMD, fallback to system 'tesseract'
                try:
                    version_output = subprocess.check_output([TESSERACT_CMD, '--version'], stderr=subprocess.STDOUT, text=True)
                    tcmd_used = TESSERACT_CMD
                except (subprocess.CalledProcessError, FileNotFoundError, OSError):
                    version_output = subprocess.check_output(['tesseract', '--version'], stderr=subprocess.STDOUT, text=True)
                    tcmd_used = 'tesseract'
                version = version_output.split('\n')[0] if version_output else "Unknown version"
                
                # Try to get available languages
                try:
                    langs_output = subprocess.check_output([tcmd_used, '--list-langs'], stderr=subprocess.STDOUT, text=True)
                except Exception:
                    langs_output = ''
                langs = [lang.strip() for lang in langs_output.split('\n')[1:] if lang.strip()]
                
                return True, {
                    "installed": True, 
                    "version": version, 
                    "languages": langs,
                    "message": "Tesseract is installed"
                }
            except (subprocess.CalledProcessError, FileNotFoundError):
                return False, {"installed": False, "message": "Tesseract is not installed or not in PATH"}
        
        else:
            return False, {"installed": False, "message": f"Unknown dependency: {name}"}
    
    except Exception as e:
        return False, {"installed": False, "message": f"Error checking {name}: {str(e)}"}

@app.before_request
def before_request():
    """Run before each request to perform housekeeping."""
    # Run cleanup periodically
    cleanup_old_files()
    
    # Make session permanent but with a timeout
    session.permanent = True

@app.route('/')
def index():
    """Home page route."""
    # Check if dependencies are properly installed
    deps_installed, message = check_dependencies()
    if not deps_installed:
        flash(message, 'error')
    return render_template('index.html')

def sanitize_text(text: Optional[str]) -> str:
    """Sanitize text by removing control characters."""
    if not text:
        return ""
    return re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F]', '', text)

def enhance_image(image: Image.Image) -> Image.Image:
    """Enhance image quality for better OCR results."""
    try:
        # Import here to avoid requiring these packages unless needed
        from PIL import ImageEnhance, ImageFilter
        
        # Apply a slight sharpening filter
        image = image.filter(ImageFilter.SHARPEN)
        
        # Increase contrast slightly
        enhancer = ImageEnhance.Contrast(image)
        image = enhancer.enhance(1.5)
        
        # Convert to grayscale if not already
        if image.mode != 'L':
            image = image.convert('L')
        
        return image
    except Exception as e:
        logger.warning(f"Image enhancement failed: {e}")
        return image  # Return original image if enhancement fails

def process_image(i: int, image_path: str, ocr_engine: str, language: str, preprocess: bool = False) -> Tuple[int, str]:
    """Process a single image with OCR (to be used in parallel processing)"""
    img_to_process = None
    try:
        text = ""
        # Open image
        logger.debug(f"Attempting to open image: {image_path}")  # Debugging log
        img_to_process = Image.open(image_path)
        logger.debug(f"Image opened successfully: {image_path}")  # Debugging log
        
        # Preprocess image if requested
        if preprocess:
            img_to_process = enhance_image(img_to_process)
        
        # Log OCR engine being used for debugging
        logger.info(f"Processing page {i+1} with OCR engine: {ocr_engine}")
        
        if ocr_engine == "tesseract":
            try:
                # Tesseract configuration for better accuracy
                config = f"--oem 1 --psm 3 -l {language}"
                if 'eng' in language and '+' not in language:
                    # Add extra parameters for English for better accuracy
                    config += " --dpi 300"
                
                # pytesseract.tesseract_cmd is set globally from TESSERACT_CMD
                try:
                    tesseract_version = pytesseract.get_tesseract_version()
                    logger.info(f"Using Tesseract version: {tesseract_version}")
                except Exception as e:
                    logger.warning(f"Could not get Tesseract version: {e}")
                
                # Log what language is being used
                logger.info(f"OCR language settings: {language}, config: {config}")
                
                text = pytesseract.image_to_string(img_to_process, config=config)
                if not text.strip():
                    logger.warning(f"Empty OCR result for page {i+1}. Trying alternative method.")
                    # Try with a different PSM mode if empty
                    config = f"--oem 1 --psm 6 -l {language}"
                    text = pytesseract.image_to_string(img_to_process, config=config)
            except Exception as e:
                logger.error(f"Tesseract OCR error: {str(e)}", exc_info=True)
                return i, f"[Error with Tesseract OCR: {str(e)}]"
        
        elif ocr_engine == "easyocr":
            try:
                import easyocr
                # Map common 3-letter ISO codes to 2-letter EasyOCR codes
                lang_map = {
                    'eng': 'en', 'fra': 'fr', 'deu': 'de', 'spa': 'es', 'ita': 'it', 'por': 'pt', 
                    'chi_sim': 'ch_sim', 'chi_tra': 'ch_tra', 'jpn': 'ja', 'kor': 'ko', 'rus': 'ru', 
                    'ara': 'ar', 'hin': 'hi'
                }
                
                # Parse and map languages (handling multiple languages separated by +)
                langs_to_load = []
                for lang in language.split('+'):
                    if lang in lang_map:
                        langs_to_load.append(lang_map[lang])
                    else:
                        langs_to_load.append(lang)
                
                # Initialize reader with all requested languages
                reader = easyocr.Reader(langs_to_load)
                
                # Process with EasyOCR (using the file path directly)
                result = reader.readtext(image_path, detail=0, paragraph=True)
                text = '\n'.join(result) if result else ""
            except Exception as e:
                logger.error(f"EasyOCR error: {str(e)}", exc_info=True)
                return i, f"[Error with EasyOCR: {str(e)}]"
        
        elif ocr_engine == "pyocr":
            try:
                import pyocr
                import pyocr.builders
                
                # Get available tools (should be Tesseract or Cuneiform)
                tools = pyocr.get_available_tools()
                if len(tools) == 0:
                    return i, "[Error: No OCR tool found for PyOCR. Install Tesseract or Cuneiform.]"
                
                # Use the first available tool (typically Tesseract)
                tool = tools[0]
                
                # Map Tesseract language codes to PyOCR
                # PyOCR uses the same language codes as Tesseract
                
                # Perform OCR
                text = tool.image_to_string(
                    img_to_process,
                    lang=language,
                    builder=pyocr.builders.TextBuilder()
                )
            except Exception as e:
                logger.error(f"PyOCR error: {str(e)}", exc_info=True)
                return i, f"[Error with PyOCR: {str(e)}]"
        
        else:
            return i, f"[Error: Unsupported OCR engine: {ocr_engine}]"

        # Sanitize text
        text = sanitize_text(text)
        
        # Attempt to detect and fix common OCR errors
        text = fix_common_ocr_errors(text)
        
        return i, text
    except FileNotFoundError as e:
        logger.error(f"File not found error processing page {i+1}: {str(e)}", exc_info=True)
        return i, f"[Error: File not found: {str(e)}. Ensure the file exists and is accessible.]"
    except Exception as e:
        logger.error(f"Error processing page {i+1} with {ocr_engine}: {str(e)}", exc_info=True)
        return i, f"[Error processing page {i+1}: {str(e)}]"
    finally:
        # Ensure PIL Image is properly closed to prevent resource leaks
        if img_to_process is not None:
            try:
                img_to_process.close()
            except Exception as e:
                logger.warning(f"Error closing image for page {i+1}: {str(e)}")
        # We don't delete the file here as it will be managed by the main process

def fix_common_ocr_errors(text: str) -> str:
    """Fix common OCR errors in text."""
    if not text:
        return text
        
    # Fix common OCR errors
    replacements = {
        # Common OCR errors
        'l1': 'h', 'rn': 'm', 'cl': 'd', 'vv': 'w',
        # Fix spaces
        ' ,': ',', ' .': '.', ' ;': ';', ' :': ':', ' !': '!', ' ?': '?',
        # Fix common misrecognitions
        '0': 'O', '1': 'I', '5': 'S',
    }
    
    # Apply replacements
    for wrong, correct in replacements.items():
        text = text.replace(wrong, correct)
    
    # Fix line breaks
    text = re.sub(r'(?<!\n)\n(?!\n)', ' ', text)  # Replace single newlines with spaces
    text = re.sub(r'\n{3,}', '\n\n', text)  # Replace multiple newlines with double newlines
    
    return text

def detect_table_in_text(text: str) -> bool:
    """Detect if text contains a potential table structure."""
    if not text or len(text.strip()) < 20:
        return False
    
    lines = text.strip().split('\n')
    if len(lines) < 3:
        return False
    
    # Look for pipe characters (markdown tables) or repeated spacing patterns
    pipe_lines = [l for l in lines if '|' in l]
    if len(pipe_lines) >= 3:
        return True
    
    # Check for aligned columns (4+ lines with consistent spacing)
    spacing_patterns = []
    for line in lines[:10]:  # Check first 10 lines
        spaces = [m.start() for m in re.finditer(r'\s{2,}', line)]
        if len(spaces) >= 2:
            spacing_patterns.append(tuple(spaces[:3]))
    
    # If we see consistent column positions, likely a table
    if len(spacing_patterns) >= 3:
        common_patterns = {}
        for pattern in spacing_patterns:
            common_patterns[pattern] = common_patterns.get(pattern, 0) + 1
        if any(count >= 2 for count in common_patterns.values()):
            return True
    
    return False

def extract_table_from_text(text: str) -> Tuple[bool, Optional[list]]:
    """Extract table data from text. Returns (is_table, table_data)."""
    lines = text.strip().split('\n')
    if len(lines) < 3:
        return False, None
    
    # Try parsing markdown table format
    if all('|' in line for line in lines[:min(3, len(lines))]):
        table_data = []
        for line in lines:
            if '|' in line:
                cells = [cell.strip() for cell in line.split('|')]
                cells = [c for c in cells if c]  # Remove empty cells
                if cells:
                    table_data.append(cells)
        if len(table_data) >= 2:
            return True, table_data
    
    # Try parsing column-aligned table (find column positions by spacing)
    column_positions = set()
    for line in lines:
        words = re.finditer(r'\S+', line)
        for match in words:
            column_positions.add(match.start())
    
    if len(column_positions) >= 2:
        column_positions = sorted(list(column_positions))
        table_data = []
        
        for line in lines:
            cells = []
            for i, pos in enumerate(column_positions):
                next_pos = column_positions[i + 1] if i + 1 < len(column_positions) else len(line)
                cell = line[pos:next_pos].strip()
                if cell:
                    cells.append(cell)
            if cells and len(cells) >= 2:
                table_data.append(cells)
        
        if len(table_data) >= 2 and all(len(row) >= 2 for row in table_data):
            return True, table_data
    
    return False, None

def format_table_as_markdown(table_data: list) -> str:
    """Format table data as markdown table."""
    if not table_data or len(table_data) < 2:
        return ""
    
    # Calculate column widths
    col_count = max(len(row) for row in table_data)
    col_widths = [0] * col_count
    
    for row in table_data:
        for i, cell in enumerate(row):
            if i < col_count:
                col_widths[i] = max(col_widths[i], len(str(cell)))
    
    # Format header
    header = table_data[0]
    header_row = ' | '.join(str(cell).ljust(col_widths[i]) for i, cell in enumerate(header))
    separator = '-|-'.join('-' * w for w in col_widths)
    
    result = header_row + '\n' + separator + '\n'
    
    # Format data rows
    for row in table_data[1:]:
        # Pad row to match column count
        padded_row = row + [''] * (col_count - len(row))
        data_row = ' | '.join(str(cell).ljust(col_widths[i]) for i, cell in enumerate(padded_row))
        result += data_row + '\n'
    
    return result

def format_table_as_html(table_data: list) -> str:
    """Format table data as HTML table."""
    if not table_data:
        return ""
    
    html = '<table border="1" cellpadding="10" cellspacing="0" style="border-collapse: collapse;">\n'
    
    # Header row
    if table_data:
        html += '  <thead>\n    <tr>\n'
        for cell in table_data[0]:
            escaped = str(cell).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            html += f'      <th>{escaped}</th>\n'
        html += '    </tr>\n  </thead>\n'
    
    # Data rows
    if len(table_data) > 1:
        html += '  <tbody>\n'
        for row in table_data[1:]:
            html += '    <tr>\n'
            for cell in row:
                escaped = str(cell).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                html += f'      <td>{escaped}</td>\n'
            html += '    </tr>\n'
        html += '  </tbody>\n'
    
    html += '</table>\n'
    return html

def format_table_in_docx(document, table_data: list) -> None:
    """Add formatted table to DOCX document."""
    if not table_data or len(table_data) < 2:
        return
    
    col_count = max(len(row) for row in table_data)
    table = document.add_table(rows=len(table_data), cols=col_count)
    table.style = 'Table Grid'
    
    # Fill in header row
    hdr_cells = table.rows[0].cells
    for i, cell_text in enumerate(table_data[0]):
        if i < len(hdr_cells):
            hdr_cells[i].text = str(cell_text)
            # Make header bold
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
    
    # Fill in data rows
    for row_idx, row_data in enumerate(table_data[1:], 1):
        if row_idx < len(table.rows):
            cells = table.rows[row_idx].cells
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < len(cells):
                    cells[col_idx].text = str(cell_text)

def process_text_with_table_detection(text: str) -> list:
    """Process text and split into segments with table detection."""
    segments = []  # List of dicts: {type: 'text'|'table', content: str or list}
    
    lines = text.split('\n')
    current_segment = []
    in_table = False
    
    for line in lines:
        # Check if line might be part of a table
        is_table_line = bool(re.match(r'^\s*\|.*\|\s*$', line)) or (
            len([m.start() for m in re.finditer(r'\s{2,}', line)]) >= 2
        )
        
        if is_table_line and not in_table:
            # Flush current text segment
            if current_segment:
                text_content = '\n'.join(current_segment).strip()
                if text_content:
                    segments.append({'type': 'text', 'content': text_content})
                current_segment = []
            in_table = True
            current_segment.append(line)
        elif not is_table_line and in_table and line.strip():
            # Line doesn't look like table, check if we should stay in table
            if not re.match(r'^\s*$', line):
                current_segment.append(line)
            else:
                # Empty line might end table
                table_text = '\n'.join(current_segment).strip()
                is_table, table_data = extract_table_from_text(table_text)
                if is_table and table_data:
                    segments.append({'type': 'table', 'content': table_data})
                else:
                    segments.append({'type': 'text', 'content': table_text})
                current_segment = []
                in_table = False
        else:
            current_segment.append(line)
    
    # Flush remaining segment
    if current_segment:
        content = '\n'.join(current_segment).strip()
        if content:
            if in_table:
                is_table, table_data = extract_table_from_text(content)
                if is_table and table_data:
                    segments.append({'type': 'table', 'content': table_data})
                else:
                    segments.append({'type': 'text', 'content': content})
            else:
                segments.append({'type': 'text', 'content': content})
    
    return segments

def save_as_markdown(text_results: Dict[int, str], output_path: str) -> None:
    """Save the extracted text results as a Markdown file with table preservation."""
    with open(output_path, 'w', encoding='utf-8') as f:
        for page_idx in sorted(text_results.keys()):
            text = text_results[page_idx]
            
            # Process text to detect tables
            segments = process_text_with_table_detection(text)
            
            for segment in segments:
                if segment['type'] == 'table':
                    # Format and write table
                    table_md = format_table_as_markdown(segment['content'])
                    f.write(table_md)
                    f.write('\n\n')
                else:
                    # Write regular text paragraphs
                    paragraphs = segment['content'].split('\n\n')
                    for para in paragraphs:
                        if para.strip():
                            f.write(para.strip() + '\n\n')
            
            # Add page separator
            if page_idx < max(text_results.keys()):
                f.write('---\n\n')


def save_as_html(text_results: Dict[int, str], output_path: str, title: str = "Converted Document") -> None:
    """Save the extracted text results as an HTML file with table preservation."""
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write('<!DOCTYPE html>\n')
        f.write('<html lang="en">\n')
        f.write('<head>\n')
        f.write('    <meta charset="UTF-8">\n')
        f.write(f'    <title>{title}</title>\n')
        f.write('    <style>\n')
        f.write('        body { font-family: Arial, sans-serif; line-height: 1.6; margin: 20px; }\n')
        f.write('        .page-break { page-break-after: always; margin: 20px 0; border-top: 2px solid #999; padding-top: 20px; }\n')
        f.write('        table { border-collapse: collapse; margin: 15px 0; width: 100%; }\n')
        f.write('        table, th, td { border: 1px solid #333; }\n')
        f.write('        th { background-color: #f2f2f2; padding: 12px; text-align: left; font-weight: bold; }\n')
        f.write('        td { padding: 10px; }\n')
        f.write('        p { margin: 10px 0; }\n')
        f.write('    </style>\n')
        f.write('</head>\n')
        f.write('<body>\n')
        f.write(f'<h1>{title}</h1>\n')
        
        for page_idx in sorted(text_results.keys()):
            text = text_results[page_idx]
            
            # Process text to detect tables
            segments = process_text_with_table_detection(text)
            
            for segment in segments:
                if segment['type'] == 'table':
                    # Format and write table
                    table_html = format_table_as_html(segment['content'])
                    f.write(table_html)
                    f.write('\n')
                else:
                    # Write regular text paragraphs
                    paragraphs = segment['content'].split('\n\n')
                    for para in paragraphs:
                        if para.strip():
                            escaped_para = para.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                            f.write(f'<p>{escaped_para.strip()}</p>\n')
            
            # Add page separator
            if page_idx < max(text_results.keys()):
                f.write('<div class="page-break"></div>\n')
                
        f.write('</body>\n')
        f.write('</html>\n')

def process_pdf_with_progress(pdf_path: str, conversion_id: str, ocr_engine: str = "tesseract", language: str = "eng", quality: str = "standard", preprocess: bool = False, orig_filename: Optional[str] = None, output_format: str = "docx") -> Tuple[bool, Optional[str], str]:
    """Process PDF with parallel processing for speed and handle different output formats"""
    temp_dir = None
    try:
        # Create a temporary directory for image files
        temp_dir = tempfile.mkdtemp(prefix="ocr_")
        logger.info(f"Created temporary directory: {temp_dir}")  # Log temp dir creation
        
        # Import PDF conversion library
        from pdf2image import convert_from_path, pdfinfo_from_path

        # Get total page count
        total_pages = 0
        try:
            pdf_info = pdfinfo_from_path(pdf_path)
            total_pages = pdf_info["Pages"]
        except Exception as e:
            logger.warning(f"Could not get page count via pdfinfo: {e}. Will count after conversion.")

        # Convert PDF to images with appropriate DPI based on quality
        dpi = 300
        if quality == "high":
            dpi = 600

        # Update status to show we're starting conversion
        if conversion_id in TASK_STATUS:
            TASK_STATUS[conversion_id].update({
                "status": "processing", 
                "step": "converting",
                "progress": 0
            })

        # Convert PDF to images but don't rely on automatic file saving
        logger.info(f"Starting PDF conversion with DPI={dpi}")
        poppler_path = r"D:\NucleusWorkSpace\PDFReaderOCR\PDFReaderOCR\Release-23.08.0-0\poppler-23.08.0\Library\bin"

        images = convert_from_path(
            pdf_path, 
            dpi=dpi,
            thread_count=1,  # Use single thread to avoid concurrency issues
            use_pdftocairo=True,  # Try to use pdftocairo which is more reliable
            fmt='png',
            poppler_path=poppler_path
            # Removed output_folder parameter to handle file saving manually
        )
        logger.info(f"PDF conversion returned {len(images)} images")

        if total_pages == 0:
            total_pages = len(images)
            
        # Manually save each image with explicit naming
        image_paths = []
        for i, img in enumerate(images):
            img_path = os.path.join(temp_dir, f'page_{i}.png')
            logger.info(f"Saving image {i+1}/{len(images)} to {img_path}")
            try:
                img.save(img_path, 'PNG')
                # Verify the file was created and has content
                if os.path.exists(img_path) and os.path.getsize(img_path) > 0:
                    image_paths.append((i, img_path))
                else:
                    logger.error(f"Failed to save image or file is empty: {img_path}")
            except Exception as e:
                logger.error(f"Error saving image {i}: {str(e)}", exc_info=True)
        
        # Check if we have all the images
        if len(image_paths) == 0:
            raise FileNotFoundError(f"No images were successfully saved from the PDF conversion")
        
        if len(image_paths) < len(images):
            logger.warning(f"Only saved {len(image_paths)} of {len(images)} images")

        logger.info(f"Successfully saved {len(image_paths)} images")

        # Update status to show conversion is complete
        if conversion_id in TASK_STATUS:
            TASK_STATUS[conversion_id].update({
                "step": "ocr",
                "progress": 10  # 10% progress after conversion
            })

        # Record start time for performance monitoring
        start_time = time.time()

        # Create a DOCX document (only if needed)
        if output_format == "docx":
            document = Document()
        
        logger.info(f"Processing PDF with 1 worker, quality: {quality}, engine: {ocr_engine}, output: {output_format}")
        
        # Process images in parallel using a process pool - using only 1 worker for reliability
        num_workers = 1 # Force single worker to avoid race conditions
        results = {}
        with ProcessPoolExecutor(max_workers=num_workers) as executor:
            # Create tasks for each image - using our saved image paths
            futures = {
                executor.submit(
                    process_image, i, img_path, ocr_engine, language, preprocess
                ): i 
                for i, img_path in image_paths
            }
            
            # Collect results as they complete
            completed = 0
            for future in as_completed(futures):
                page_idx, text = future.result()
                results[page_idx] = text
                
                # Update progress (allocate 10-90% for OCR process)
                completed += 1
                progress = 10 + int((completed / len(image_paths)) * 80)
                if conversion_id in TASK_STATUS:
                    TASK_STATUS[conversion_id]["progress"] = progress
        
        # Update status to show we're assembling the document
        if conversion_id in TASK_STATUS:
            TASK_STATUS[conversion_id].update({
                "step": "assembling",
                "progress": 90
            })
            
        # Determine output filename and path
        document_name = orig_filename or 'document.pdf'
        base_filename = os.path.splitext(document_name)[0]
        output_filename = f"{secure_clean_filename(base_filename)}.{output_format}"
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{conversion_id}_{output_filename}")

        # Assemble and save the document based on the format
        if output_format == "docx":
            for i in range(len(image_paths)):
                if i in results:
                    text = results[i]
                    
                    # Process text to detect tables
                    segments = process_text_with_table_detection(text)
                    
                    for segment in segments:
                        if segment['type'] == 'table':
                            # Add table to document
                            format_table_in_docx(document, segment['content'])
                        else:
                            # Add text paragraph
                            if segment['content'].strip():
                                document.add_paragraph(segment['content'].strip())
                    
                    if i < len(image_paths) - 1:
                        document.add_page_break()
            document.save(output_path)
        elif output_format == "txt":
            with open(output_path, 'w', encoding='utf-8') as f:
                for i in sorted(results.keys()):
                    text = results[i]
                    
                    # Process text to detect tables
                    segments = process_text_with_table_detection(text)
                    
                    for segment in segments:
                        if segment['type'] == 'table':
                            # Format table as simple text with alignment
                            table_data = segment['content']
                            if table_data:
                                # Calculate column widths
                                col_widths = [0] * max(len(row) for row in table_data)
                                for row in table_data:
                                    for j, cell in enumerate(row):
                                        col_widths[j] = max(col_widths[j], len(str(cell)))
                                
                                # Write header
                                header = table_data[0]
                                f.write(' | '.join(str(c).ljust(col_widths[j]) for j, c in enumerate(header)) + '\n')
                                f.write('-+-'.join('-' * w for w in col_widths) + '\n')
                                
                                # Write data rows
                                for row in table_data[1:]:
                                    padded = row + [''] * (len(col_widths) - len(row))
                                    f.write(' | '.join(str(c).ljust(col_widths[j]) for j, c in enumerate(padded)) + '\n')
                                f.write('\n')
                        else:
                            # Write regular text
                            if segment['content'].strip():
                                f.write(segment['content'].strip() + '\n')
                    
                    # Add a separator between pages for clarity
                    if i < max(results.keys()):
                        f.write("\n--- Page Break ---\n\n")
        elif output_format == "md":
            save_as_markdown(results, output_path)
        elif output_format == "html":
            save_as_html(results, output_path, title=base_filename)
        else:
            raise ValueError(f"Unsupported output format: {output_format}")

        # Log performance metrics
        elapsed_time = time.time() - start_time
        pages_per_second = total_pages / elapsed_time if elapsed_time > 0 else 0
        logger.info(f"PDF processing completed. Pages: {total_pages}, Time: {elapsed_time:.2f}s, Pages/sec: {pages_per_second:.2f}, Output: {output_path}")

        # Clean up original PDF after successful conversion
        if os.path.exists(pdf_path):
            os.remove(pdf_path)

        # Update status to complete
        if conversion_id in TASK_STATUS:
            TASK_STATUS[conversion_id]["progress"] = 100

        return True, output_path, output_filename # Return the actual path and filename

    except ImportError as e:
        # Specific handling for missing OCR engine imports
        error_message = f"Error: The required OCR engine '{ocr_engine}' is not properly installed. Please install it with pip: {str(e)}"
        logger.error(error_message)
        return False, None, error_message
    except Exception as e:
        logger.error(f"Error during PDF processing: {str(e)}", exc_info=True)
        error_message = f"An unexpected error occurred: {str(e)}"
        return False, None, error_message
    finally:
        # Clean up temporary directory
        if temp_dir and os.path.exists(temp_dir):
            try:
                shutil.rmtree(temp_dir)
                logger.info(f"Cleaned up temporary directory: {temp_dir}") # Log temp dir cleanup
            except Exception as e:
                logger.error(f"Error cleaning up temporary directory: {e}")

def run_task_in_background(func: callable, task_id: str, *args: Any, **kwargs: Any) -> str:
    """Run a function in a background thread and track its status"""
    def task_wrapper():
        try:
            # Initialize task status with timestamp for cleanup
            TASK_STATUS[task_id] = {
                "status": "processing", 
                "step": "initializing",
                "progress": 0, 
                "timestamp": time.time()
            }
            
            # Run the actual task
            result = func(*args, **kwargs)
            
            # Store results and update status
            TASK_RESULTS[task_id] = result
            TASK_STATUS[task_id].update({
                "status": "completed", 
                "progress": 100,
                "timestamp": time.time()  # Update timestamp
            })
            
        except Exception as e:
            logger.error(f"Background task error: {str(e)}", exc_info=True)
            TASK_STATUS[task_id] = {
                "status": "failed", 
                "error": str(e), 
                "progress": 0,
                "timestamp": time.time()
            }
    
    thread = Thread(target=task_wrapper)
    thread.daemon = True
    thread.start()
    return task_id

@app.route('/upload', methods=['POST'])
def upload_file():
    """Handle file upload and initiate OCR processing."""
    try:
        # Check dependencies first
        deps_installed, message = check_dependencies()
        if not deps_installed:
            flash(message, 'error')
            flash("Please install the required dependencies before proceeding", 'error')
            return redirect(url_for('index'))

        if 'file' not in request.files:
            flash('No file part', 'error')
            return redirect(url_for('index'))

        file = request.files['file']
        if file.filename == '':
            flash('No selected file', 'error')
            return redirect(url_for('index'))

        if not allowed_file(file.filename):
            flash('Invalid file type. Please upload a PDF.', 'error')
            return redirect(url_for('index'))

        # Generate unique ID for this conversion
        conversion_id = str(uuid.uuid4())
        session['conversion_id'] = conversion_id

        # Save original filename for later use (but sanitize it)
        orig_filename = secure_clean_filename(file.filename)
        session['orig_filename'] = orig_filename

        # Get OCR options from form
        ocr_engine = request.form.get('ocr-engine', 'tesseract')
        language = request.form.get('language', 'eng')
        quality = request.form.get('ocr-quality', 'standard')
        preprocess = request.form.get('preprocess', '0') == '1'
        output_format = request.form.get('output-format', 'docx')

        # Log processing request
        logger.info(f"Processing request: file={orig_filename}, engine={ocr_engine}, lang={language}, quality={quality}, preprocess={preprocess}, format={output_format}")

        # Create a temporary filename to avoid collisions
        temp_filename = f"{conversion_id}_{secure_clean_filename(file.filename)}"
        pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], temp_filename)

        # Save the uploaded file
        try:
            file.save(pdf_path)
            logger.info(f"Saved uploaded file to {pdf_path}")
        except Exception as e:
            logger.error(f"Error saving uploaded file: {str(e)}", exc_info=True)
            flash("An error occurred while saving the uploaded file. Please try again.", 'error')
            return redirect(url_for('index'))

        session['pdf_path'] = pdf_path

        # Process asynchronously
        try:
            task_id = run_task_in_background(
                process_pdf_with_progress,
                conversion_id,
                pdf_path,
                conversion_id,
                ocr_engine,
                language,
                quality,
                preprocess,
                orig_filename,
                output_format
            )
            session['task_id'] = conversion_id
            return redirect(url_for('status', task_id=conversion_id))
        except Exception as e:
            logger.error(f"Error starting background task: {str(e)}", exc_info=True)
            flash("An error occurred while starting the OCR process. Please try again.", 'error')
            return redirect(url_for('index'))

    except Exception as e:
        logger.error(f"Unexpected error during file upload: {str(e)}", exc_info=True)
        flash("An unexpected error occurred. Please try again.", 'error')
        return redirect(url_for('index'))

@app.route('/status/<task_id>')
def status(task_id):
    """Display status page for an ongoing conversion."""
    # Check if task exists
    if task_id not in TASK_STATUS:
        flash("The requested conversion was not found.", 'error')
        return redirect(url_for('index'))
        
    return render_template('status.html', task_id=task_id)

@app.route('/api/task_status/<task_id>')
def task_status(task_id):
    """API endpoint to check task status"""
    if task_id in TASK_STATUS:
        status_data = TASK_STATUS[task_id].copy()
        
        # If task is completed, include the path to results
        if status_data.get("status") == "completed" and task_id in TASK_RESULTS:
            success, result_path, output_filename = TASK_RESULTS[task_id]
            if success:
                # Store results in session
                session['result_path'] = result_path # Use generic name
                session['output_filename'] = output_filename
                status_data["redirect"] = url_for('success')
            else:
                # Store error
                status_data["error"] = output_filename
                status_data["redirect"] = url_for('index')
                flash(f"Conversion failed: {output_filename}", 'error')
        
        return jsonify(status_data)
    
    return jsonify({"status": "not_found"})

@app.route('/success')
def success():
    """Display success page after successful conversion."""
    # Check if we have valid session data
    if 'result_path' not in session or 'output_filename' not in session: # Updated session key
        flash('No conversion data found. Please upload a file first.', 'error')
        return redirect(url_for('index'))

    return render_template('success.html', filename=session['output_filename'])

@app.route('/download')
def download_file():
    """Provide the converted file for download."""
    # Check if we have valid session data
    if 'result_path' not in session or 'output_filename' not in session: # Updated session key
        flash('No conversion data found. Please upload a file first.', 'error')
        return redirect(url_for('index'))

    result_path = session['result_path'] # Updated session key
    output_filename = session['output_filename']

    # Security check - ensure file exists and is within uploads directory
    if not os.path.exists(result_path) or not os.path.isfile(result_path):
        flash('The converted file is no longer available.', 'error')
        return redirect(url_for('index'))
    
    if not os.path.abspath(result_path).startswith(os.path.abspath(UPLOAD_FOLDER)):
        logger.error(f"Security issue: Attempted to access file outside uploads directory: {result_path}")
        flash('Access denied.', 'error')
        return redirect(url_for('index'))

    # Determine MIME type based on file extension
    _, ext = os.path.splitext(output_filename)
    ext = ext.lower()
    mime_types = {
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.txt': 'text/plain',
        '.md': 'text/markdown',
        '.html': 'text/html',
    }
    mime_type = mime_types.get(ext, 'application/octet-stream') # Default MIME type

    try:
        return send_file(
            result_path,
            mimetype=mime_type, # Use determined MIME type
            as_attachment=True,
            download_name=output_filename
        )
    except Exception as e:
        logger.error(f"Error during file download: {str(e)}", exc_info=True)
        flash(f"Error downloading file: {str(e)}", 'error')
        return redirect(url_for('index'))

@app.route('/new_conversion')
def new_conversion():
    """Start a new conversion, cleaning up any existing files."""
    # Clean up session data and files
    pdf_path = session.pop('pdf_path', None)
    if (pdf_path and os.path.exists(pdf_path)):
        try:
            os.remove(pdf_path)
            logger.info(f"Cleaned up PDF file: {pdf_path}")
        except OSError as e:
            logger.warning(f"Could not remove PDF file {pdf_path}: {e}")

    result_path = session.pop('result_path', None) # Updated session key
    if (result_path and os.path.exists(result_path)):
         try:
            os.remove(result_path)
            logger.info(f"Cleaned up result file: {result_path}")
         except OSError as e:
            logger.warning(f"Could not remove result file {result_path}: {e}")

    # Clear remaining session data
    session.pop('conversion_id', None)
    session.pop('orig_filename', None)
    session.pop('output_filename', None)

    # Redirect to index
    return redirect(url_for('index'))

@app.route('/api/check-dependency')
def api_check_dependency():
    """API endpoint to check if a specific dependency is installed"""
    name = request.args.get('name', '')
    if not name:
        return jsonify({"error": "No dependency name provided"}), 400
    
    installed, data = check_dependency(name)
    return jsonify(data)

@app.route('/system-check')
def system_check():
    """Check system status and dependencies"""
    results = {
        "status": "ok",
        "errors": [],
        "dependencies": {}
    }
    
    # Check Python version
    results["python_version"] = sys.version
    
    # Check Tesseract
    try:
        tesseract_installed, tesseract_data = check_dependency('tesseract')
        results["dependencies"]["tesseract"] = tesseract_data
        
        if not tesseract_installed:
            results["status"] = "error"
            results["errors"].append("Tesseract OCR is not installed or not found in PATH")
    except Exception as e:
        results["dependencies"]["tesseract"] = {"error": str(e)}
        results["status"] = "error"
        results["errors"].append(f"Error checking Tesseract: {str(e)}")
    
    # Check Poppler
    try:
        poppler_installed, poppler_data = check_dependency('poppler')
        results["dependencies"]["poppler"] = poppler_data
        
        if not poppler_installed:
            results["status"] = "error"
            results["errors"].append("Poppler is not installed or not found in PATH")
    except Exception as e:
        results["dependencies"]["poppler"] = {"error": str(e)}
        results["status"] = "error"
        results["errors"].append(f"Error checking Poppler: {str(e)}")
    
    # Check upload directory
    upload_dir = app.config['UPLOAD_FOLDER']
    results["upload_dir"] = {
        "path": upload_dir,
        "exists": os.path.exists(upload_dir),
        "writable": os.access(upload_dir, os.W_OK) if os.path.exists(upload_dir) else False
    }
    
    if not results["upload_dir"]["exists"]:
        results["status"] = "error"
        results["errors"].append(f"Upload directory does not exist: {upload_dir}")
    elif not results["upload_dir"]["writable"]:
        results["status"] = "error"
        results["errors"].append(f"Upload directory is not writable: {upload_dir}")
    
    return jsonify(results)

@app.errorhandler(404)
def page_not_found(e):
    """Handle 404 errors."""
    return render_template('error.html', error="Page not found", code=404), 404

@app.errorhandler(500)
def server_error(e):
    """Handle 500 errors."""
    logger.error(f"Server error: {str(e)}", exc_info=True)
    return render_template('error.html', error="Internal server error", code=500), 500

@app.errorhandler(413)
def request_entity_too_large(e):
    """Handle file too large errors."""
    flash('The file is too large. Maximum size is 64MB.', 'error')
    return redirect(url_for('index'))

if __name__ == '__main__':
    # Set the start method for multiprocessing
    multiprocessing.set_start_method('spawn', force=True)
    
    # Initial dependency check
    deps_installed, message = check_dependencies()
    if not deps_installed:
        logger.warning(f"Dependency issue: {message}")
        print(f"Warning: {message}")
        print("The application will start but may not work correctly until all dependencies are installed.")

    # Server configuration
    host = '0.0.0.0'
    port = int(os.environ.get('PORT', 8011))
    debug = os.environ.get('FLASK_ENV') == 'development'

    app.run(debug=debug, host=host, port=port)